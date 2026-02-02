"""
Test script for OCR Service
Tests resume text extraction from PDF and DOCX files.
"""

import sys
from pathlib import Path
import os

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from services.ocr_service import OCRService


def interactive_demo():
    """Interactive demo - upload your own file and see extracted text."""
    print("\n" + "=" * 70)
    print("📄 OCR SERVICE - INTERACTIVE DEMO")
    print("=" * 70)
    print("\nThis demo lets you upload your own PDF or DOCX file")
    print("and see the extracted text and parsed resume data.")
    print("=" * 70)
    
    service = OCRService()
    
    print(f"\n✅ OCR Service initialized")
    print(f"  - PDF Support: {'✓' if service.has_pdf else '✗'}")
    print(f"  - DOCX Support: {'✓' if service.has_docx else '✗'}")
    print(f"  - OCR Support: {'✓' if service.has_ocr else '✗'}")
    
    if not service.has_pdf and not service.has_docx:
        print("\n❌ Error: No document libraries installed!")
        print("Install with: pip install pdfplumber python-docx")
        return
    
    print("\n" + "=" * 70)
    print("📁 UPLOAD YOUR FILE")
    print("=" * 70)
    print("\nEnter the full path to your PDF or DOCX file.")
    print("Example: C:\\Users\\YourName\\Documents\\resume.pdf")
    print("Or drag and drop the file here and press Enter.")
    print("\nPress Enter without typing anything to exit.")
    print("-" * 70)
    
    file_path = input("\nFile path: ").strip()
    
    # Remove quotes if user dragged and dropped
    file_path = file_path.strip('"').strip("'")
    
    if not file_path:
        print("\n👋 Exiting demo...")
        return
    
    # Convert to Path object
    file_path = Path(file_path)
    
    # Check if file exists
    if not file_path.exists():
        print(f"\n❌ Error: File not found: {file_path}")
        print("Please check the path and try again.")
        return
    
    # Check file extension
    file_ext = file_path.suffix.lower()
    if file_ext not in ['.pdf', '.docx', '.doc']:
        print(f"\n❌ Error: Unsupported file format: {file_ext}")
        print("Supported formats: .pdf, .docx, .doc")
        return
    
    print(f"\n✅ File found: {file_path.name}")
    print(f"📊 File size: {file_path.stat().st_size / 1024:.2f} KB")
    print(f"📝 Format: {file_ext.upper()}")
    
    print("\n" + "=" * 70)
    print("🔍 EXTRACTING TEXT...")
    print("=" * 70)
    
    try:
        # Extract text
        text = service.extract_text(file_path)
        
        print(f"\n✅ Successfully extracted {len(text)} characters!")
        
        # Show preview
        print("\n" + "=" * 70)
        print("📄 TEXT PREVIEW (First 500 characters)")
        print("=" * 70)
        print(text[:500])
        if len(text) > 500:
            print("\n... (truncated)")
        print("=" * 70)
        
        # Parse resume
        print("\n" + "=" * 70)
        print("🔍 PARSING RESUME DATA...")
        print("=" * 70)
        
        parsed = service.parse_resume(file_path)
        
        print("\n✅ Resume parsed successfully!")
        print("\n" + "=" * 70)
        print("📊 EXTRACTED INFORMATION")
        print("=" * 70)
        
        # Email
        if parsed['email']:
            print(f"\n📧 Email: {parsed['email']}")
        else:
            print(f"\n📧 Email: Not found")
        
        # Phone
        if parsed['phone']:
            print(f"📞 Phone: {parsed['phone']}")
        else:
            print(f"📞 Phone: Not found")
        
        # Skills
        print(f"\n💼 Skills Found: {len(parsed['skills'])}")
        if parsed['skills']:
            print("   Skills:")
            for i, skill in enumerate(parsed['skills'][:10], 1):
                print(f"   {i}. {skill}")
            if len(parsed['skills']) > 10:
                print(f"   ... and {len(parsed['skills']) - 10} more")
        else:
            print("   No technical skills detected")
        
        # Sections
        print(f"\n📑 Sections Identified: {len(parsed['sections'])}")
        if parsed['sections']:
            for section in parsed['sections'].keys():
                print(f"   • {section.title()}")
        
        # Full text option
        print("\n" + "=" * 70)
        show_full = input("\n📄 Show full extracted text? (y/n): ").strip().lower()
        if show_full == 'y':
            print("\n" + "=" * 70)
            print("📄 FULL EXTRACTED TEXT")
            print("=" * 70)
            print(text)
            print("=" * 70)
        
        # Save option
        print("\n" + "=" * 70)
        save_text = input("\n💾 Save extracted text to file? (y/n): ").strip().lower()
        if save_text == 'y':
            output_file = file_path.parent / f"{file_path.stem}_extracted.txt"
            output_file.write_text(text, encoding='utf-8')
            print(f"\n✅ Text saved to: {output_file}")
        
        print("\n" + "=" * 70)
        print("✅ DEMO COMPLETED SUCCESSFULLY!")
        print("=" * 70)
        
    except Exception as e:
        print(f"\n❌ Error processing file: {e}")
        import traceback
        traceback.print_exc()


def create_sample_docx():
    """Create a sample DOCX resume for testing."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Skipping DOCX test.")
        return None
    
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Email: john.doe@example.com | Phone: (555) 123-4567')
    
    doc.add_heading('Education', 1)
    doc.add_paragraph('B.S. Computer Science, University of Technology, 2020')
    
    doc.add_heading('Skills', 1)
    doc.add_paragraph('Python, JavaScript, React, Node.js, FastAPI, MongoDB, Docker, AWS, Machine Learning')
    
    doc.add_heading('Experience', 1)
    doc.add_paragraph('Senior Software Engineer at Tech Corp (2020-Present)')
    doc.add_paragraph('• Developed REST APIs using FastAPI and Python')
    doc.add_paragraph('• Implemented microservices architecture with Docker and Kubernetes')
    
    doc.add_heading('Projects', 1)
    doc.add_paragraph('AI Interview Simulator - Built using Python, FastAPI, OpenAI GPT')
    
    test_file = Path(__file__).parent / "test_resume.docx"
    doc.save(test_file)
    return test_file


def create_sample_pdf():
    """Create a sample PDF resume for testing."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        print("reportlab not installed. Skipping PDF test.")
        print("Install with: pip install reportlab")
        return None
    
    test_file = Path(__file__).parent / "test_resume.pdf"
    c = canvas.Canvas(str(test_file), pagesize=letter)
    
    # Title
    c.setFont("Helvetica-Bold", 20)
    c.drawString(100, 750, "Jane Smith")
    
    # Contact
    c.setFont("Helvetica", 12)
    c.drawString(100, 730, "Email: jane.smith@example.com | Phone: (555) 987-6543")
    
    # Education
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 700, "Education")
    c.setFont("Helvetica", 12)
    c.drawString(100, 680, "M.S. Artificial Intelligence, Tech University, 2021")
    
    # Skills
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 650, "Skills")
    c.setFont("Helvetica", 12)
    c.drawString(100, 630, "Python, TensorFlow, PyTorch, Deep Learning, NLP, Computer Vision")
    c.drawString(100, 615, "SQL, PostgreSQL, MongoDB, AWS, Docker, Kubernetes")
    
    # Experience
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 585, "Experience")
    c.setFont("Helvetica", 12)
    c.drawString(100, 565, "ML Engineer at AI Startup (2021-Present)")
    c.drawString(100, 550, "• Built NLP models using transformers and BERT")
    c.drawString(100, 535, "• Deployed models to production using FastAPI and Docker")
    
    c.save()
    return test_file


def test_ocr_basic():
    """Test basic OCR service functionality."""
    print("\n" + "=" * 70)
    print("TEST 1: Basic OCR Service Initialization")
    print("=" * 70)
    
    service = OCRService()
    print(f"✓ Service initialized")
    print(f"  - PDF Support: {service.has_pdf}")
    print(f"  - DOCX Support: {service.has_docx}")
    print(f"  - OCR Support: {service.has_ocr}")
    print(f"  - Supported formats: {service.supported_formats}")


def test_docx_extraction():
    """Test DOCX text extraction."""
    print("\n" + "=" * 70)
    print("TEST 2: DOCX Text Extraction")
    print("=" * 70)
    
    service = OCRService()
    
    if not service.has_docx:
        print("⚠ SKIPPED: python-docx not installed")
        return
    
    # Create sample DOCX
    test_file = create_sample_docx()
    if not test_file:
        print("⚠ SKIPPED: Could not create test file")
        return
    
    try:
        print(f"Created test file: {test_file.name}")
        
        # Extract text
        text = service.extract_text(test_file)
        print(f"✓ Extracted {len(text)} characters")
        print(f"\nFirst 200 characters:")
        print("-" * 70)
        print(text[:200])
        print("-" * 70)
        
        # Parse resume
        parsed = service.parse_resume(test_file)
        print(f"\n✓ Parsed resume data:")
        print(f"  - Email: {parsed['email']}")
        print(f"  - Phone: {parsed['phone']}")
        print(f"  - Skills found: {len(parsed['skills'])}")
        print(f"  - Skills: {', '.join(parsed['skills'][:5])}...")
        print(f"  - Sections: {list(parsed['sections'].keys())}")
        
    finally:
        # Cleanup
        if test_file and test_file.exists():
            test_file.unlink()
            print(f"\n✓ Cleaned up test file")


def test_pdf_extraction():
    """Test PDF text extraction."""
    print("\n" + "=" * 70)
    print("TEST 3: PDF Text Extraction")
    print("=" * 70)
    
    service = OCRService()
    
    if not service.has_pdf:
        print("⚠ SKIPPED: pdfplumber not installed")
        return
    
    # Create sample PDF
    test_file = create_sample_pdf()
    if not test_file:
        print("⚠ SKIPPED: Could not create test file")
        return
    
    try:
        print(f"Created test file: {test_file.name}")
        
        # Extract text
        text = service.extract_text(test_file)
        print(f"✓ Extracted {len(text)} characters")
        print(f"\nFirst 200 characters:")
        print("-" * 70)
        print(text[:200])
        print("-" * 70)
        
        # Parse resume
        parsed = service.parse_resume(test_file)
        print(f"\n✓ Parsed resume data:")
        print(f"  - Email: {parsed['email']}")
        print(f"  - Phone: {parsed['phone']}")
        print(f"  - Skills found: {len(parsed['skills'])}")
        print(f"  - Skills: {', '.join(parsed['skills'][:5])}...")
        
    finally:
        # Cleanup
        if test_file and test_file.exists():
            test_file.unlink()
            print(f"\n✓ Cleaned up test file")


def test_error_handling():
    """Test error handling."""
    print("\n" + "=" * 70)
    print("TEST 4: Error Handling")
    print("=" * 70)
    
    service = OCRService()
    
    # Test 1: Non-existent file
    try:
        service.extract_text("nonexistent_file.pdf")
        print("✗ Should have raised FileNotFoundError")
    except FileNotFoundError as e:
        print(f"✓ Correctly raised FileNotFoundError: {e}")
    
    # Test 2: Unsupported format
    test_file = Path(__file__).parent / "test.txt"
    test_file.write_text("Test content")
    
    try:
        service.extract_text(test_file)
        print("✗ Should have raised ValueError for unsupported format")
    except ValueError as e:
        print(f"✓ Correctly raised ValueError: {e}")
    finally:
        test_file.unlink()


def run_all_tests():
    """Run all OCR tests."""
    print("\n" + "=" * 70)
    print(" OCR SERVICE TEST SUITE")
    print("=" * 70)
    
    test_ocr_basic()
    test_docx_extraction()
    test_pdf_extraction()
    test_error_handling()
    
    print("\n" + "=" * 70)
    print(" ALL TESTS COMPLETED")
    print("=" * 70)
    print("\nNote: If some tests were skipped, install the required dependencies:")
    print("  pip install pdfplumber python-docx reportlab")


def main_menu():
    """Main menu for test selection."""
    while True:
        print("\n" + "=" * 70)
        print("📄 OCR SERVICE TEST MENU")
        print("=" * 70)
        print("\n  1. Interactive Demo (Upload Your Own File)")
        print("  2. Run Automated Tests")
        print("  3. Exit")
        print("\n" + "=" * 70)
        
        choice = input("\nSelect option (1-3): ").strip()
        
        if choice == '1':
            interactive_demo()
        elif choice == '2':
            run_all_tests()
        elif choice == '3':
            print("\n👋 Goodbye!")
            break
        else:
            print("\n❌ Invalid option. Please select 1-3.")


if __name__ == "__main__":
    main_menu()

