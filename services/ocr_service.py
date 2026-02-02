"""
OCR Service for Marco AI Interview Simulator
Extracts text from PDF and DOCX resume files with OCR support for scanned documents.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Union
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class OCRService:
    """
    Service for extracting and parsing text from resume files.
    Supports PDF, DOCX, and scanned documents via OCR.
    """
    
    def __init__(self):
        """Initialize OCR Service with required dependencies check."""
        self._check_dependencies()
        self.supported_formats = ['.pdf', '.docx', '.doc']
        
    def _check_dependencies(self):
        """Check if required libraries are available."""
        self.has_pdf = False
        self.has_docx = False
        self.has_ocr = False
        self.ocr_reader = None
        
        try:
            import pdfplumber
            self.has_pdf = True
        except ImportError:
            logger.warning("pdfplumber not installed. PDF support disabled.")
            
        try:
            import docx
            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX support disabled.")
            
        try:
            import easyocr
            self.ocr_reader = easyocr.Reader(['en'], gpu=False)  # CPU mode for compatibility
            self.has_ocr = True
            logger.info("EasyOCR initialized successfully")
        except ImportError:
            logger.warning("easyocr not installed. OCR support disabled.")
        except Exception as e:
            logger.warning(f"EasyOCR initialization failed: {e}. OCR support disabled.")
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a resume file.
        
        Args:
            file_path: Path to the resume file (PDF or DOCX)
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Extracting text from: {file_path.name}")
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Format {file_ext} not implemented")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        if not self.has_pdf:
            raise RuntimeError("pdfplumber not installed. Install with: pip install pdfplumber")
        
        import pdfplumber
        
        text_content = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_content.append(page_text)
                    else:
                        # If no text found, try OCR
                        logger.info(f"Page {page_num} has no extractable text, attempting OCR...")
                        ocr_text = self._ocr_page(page)
                        if ocr_text:
                            text_content.append(ocr_text)
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise RuntimeError(f"Failed to extract PDF: {e}")
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        if not self.has_docx:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        import docx
        
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise RuntimeError(f"Failed to extract DOCX: {e}")
    
    def _ocr_page(self, page) -> str:
        """Perform OCR on a PDF page using EasyOCR."""
        if not self.has_ocr or self.ocr_reader is None:
            logger.warning("OCR not available. Install easyocr: pip install easyocr")
            return ""
        
        try:
            from PIL import Image
            import io
            import numpy as np
            
            # Convert page to image
            img = page.to_image(resolution=300)
            pil_img = img.original
            
            # Convert PIL image to numpy array for EasyOCR
            img_array = np.array(pil_img)
            
            # Perform OCR
            results = self.ocr_reader.readtext(img_array)
            
            # Extract text from results
            text = ' '.join([result[1] for result in results])
            return text
            
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""
    
    def parse_resume(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Extract and parse resume information.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary containing parsed resume data:
            - raw_text: Full extracted text
            - skills: List of identified skills
            - email: Extracted email address
            - phone: Extracted phone number
            - sections: Dictionary of identified sections
        """
        text = self.extract_text(file_path)
        
        parsed_data = {
            'raw_text': text,
            'skills': self._extract_skills(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'sections': self._identify_sections(text)
        }
        
        logger.info(f"Parsed resume: Found {len(parsed_data['skills'])} skills")
        return parsed_data
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        # Common technical skills keywords
        skill_keywords = [
            # Programming languages
            'python', 'java', 'javascript', 'c\\+\\+', 'c#', 'ruby', 'php', 'swift',
            'kotlin', 'go', 'rust', 'typescript', 'scala', 'r\\b',
            # Web technologies
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 
            'django', 'flask', 'fastapi', 'spring', 'asp.net',
            # Databases
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 
            'sqlite', 'cassandra', 'dynamodb',
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
            'ci/cd', 'terraform', 'ansible',
            # AI/ML
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 
            'scikit-learn', 'nlp', 'computer vision', 'opencv',
            # Other
            'rest api', 'graphql', 'microservices', 'agile', 'scrum'
        ]
        
        found_skills = []
        text_lower = text.lower()
        
        for skill in skill_keywords:
            # Use word boundary matching
            pattern = r'\b' + skill + r'\b'
            if re.search(pattern, text_lower, re.IGNORECASE):
                # Store original case version if found
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    found_skills.append(match.group())
        
        # Remove duplicates and return
        return list(set(found_skills))
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group() if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        # Pattern for common phone formats
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # +1-234-567-8900
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # (234) 567-8900
            r'\d{10}',  # 2345678900
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group()
        
        return None
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify common resume sections."""
        sections = {}
        
        # Common section headers
        section_patterns = {
            'education': r'(?i)(education|academic|qualification)',
            'experience': r'(?i)(experience|employment|work history)',
            'skills': r'(?i)(skills|technical skills|competencies)',
            'projects': r'(?i)(projects|portfolio)',
            'certifications': r'(?i)(certifications|certificates)'
        }
        
        lines = text.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            # Check if line is a section header
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line):
                    current_section = section_name
                    sections[section_name] = line
                    break
        
        return sections


# Standalone testing function
def test_ocr_service():
    """Test the OCR service with a sample file."""
    service = OCRService()
    
    print("=" * 60)
    print("OCR Service Test")
    print("=" * 60)
    print(f"PDF Support: {service.has_pdf}")
    print(f"DOCX Support: {service.has_docx}")
    print(f"OCR Support: {service.has_ocr}")
    print(f"Supported formats: {service.supported_formats}")
    print("=" * 60)
    
    # Create a test text file to demonstrate
    test_file = Path("test_resume.txt")
    test_content = """
John Doe
Email: john.doe@email.com
Phone: (555) 123-4567

EDUCATION
B.S. Computer Science, University of Technology

SKILLS
Python, JavaScript, React, Node.js, MongoDB, AWS, Docker, Machine Learning

EXPERIENCE
Senior Software Engineer at Tech Corp
- Developed REST APIs using FastAPI and Python
- Implemented microservices architecture with Docker and Kubernetes

PROJECTS
AI Interview Simulator - Built using Python, FastAPI, and OpenAI GPT
    """
    
    test_file.write_text(test_content)
    print(f"\nCreated test file: {test_file}")
    
    try:
        # Note: This will fail because .txt is not supported, demonstrating error handling
        result = service.extract_text(test_file)
        print(f"\nExtracted text:\n{result[:200]}...")
    except ValueError as e:
        print(f"\nExpected error (testing validation): {e}")
    finally:
        test_file.unlink()
        print("\nTest file cleaned up.")


if __name__ == "__main__":
    test_ocr_service()
